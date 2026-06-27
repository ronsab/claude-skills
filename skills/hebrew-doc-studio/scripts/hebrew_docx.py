#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
hebrew_docx.py - בונה מסמכי Word (.docx) בעברית עם RTL מלא וטיפול נכון בטקסט מעורב.

הכלי הזה פותר את שלוש התקלות שמשבשות מסמכי Word בעברית:
  1. בסיס הפסקה לא RTL  -> שורות נשפכות שמאל-לימין.
  2. כל השורה ב-run אחד מסומן rtl כולל אנגלית  -> האנגלית/מספרים קופצים לצד הלא נכון.
  3. הגדרת font/size רק על w:ascii/w:sz בלי w:cs/w:szCs -> הגדרות הגופן פשוט לא חלות על העברית.

עקרון מנחה: ב-Word *לא* מריצים get_display ולא מסדרים תווים ידנית. Word מריץ את
אלגוריתם ה-bidi בעצמו בזמן הרינדור. התפקיד שלנו הוא רק לפלוט את דגלי ה-XML הנכונים.

תלות: python-docx בלבד. (אין תלות בגופנים בזמן יצירה - Word מרנדר אצל המשתמש.)

שימוש מהיר:
    from hebrew_docx import HebrewDoc
    doc = HebrewDoc(font="Assistant", size=12)
    doc.heading("חוזה שירותים", level=0)
    doc.paragraph('ההסכם נחתם בין חברת Acme בע"מ לבין הלקוח (גרסה 2).')
    doc.bullets(["סעיף ראשון", "סעיף שני עם מונח SaaS באנגלית"])
    doc.table(
        headers=["תיאור", "כמות", 'מחיר (ש"ח)'],
        rows=[["ייעוץ", "3", "1,500.00"], ["פיתוח", "1", "12,000.00"]],
    )
    doc.save("contract.docx")
"""

import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("חסרה תלות. התקן עם:  pip install --break-system-packages python-docx", file=sys.stderr)
    raise

# טווח עברי + צורות הצגה עבריות. משמש לבחירת כיוון לכל run.
_HEB = re.compile(r'[\u0590-\u05FF\uFB1D-\uFB4F]')


def _strong(ch):
    """True לאות עברית, False לאות לטינית חזקה, None לתו ניטרלי (רווח/ספרה/פיסוק)."""
    if _HEB.match(ch):
        return True
    if ch.isalpha():
        return False
    return None


def split_by_script(text):
    """מפצל מחרוזת מעורבת לרצף (segment, is_rtl).

    כיוון כל קטע נקבע לפי התווים *החזקים* שבו; תווים ניטרליים (רווחים, ספרות,
    פיסוק) נדבקים לקטע הנוכחי. ניטרלים בתחילת המחרוזת יורשים את כיוון התו החזק
    הראשון בכל המחרוזת (ברירת מחדל RTL למחרוזת ניטרלית לגמרי במסמך עברי), כדי
    ששורה לטינית שמתחילה בספרה/סוגריים לא תסומן בטעות RTL. Word מתקן את הסדר
    החזותי הסופי, אז הפיצול רק צריך להכניס את התווים החזקים ל-runs מסומנים נכון.
    """
    default_rtl = next((s for s in (_strong(c) for c in text) if s is not None), True)
    segments, buf, buf_rtl = [], '', None
    for ch in text:
        s = _strong(ch)
        kind = s if s is not None else (buf_rtl if buf_rtl is not None else default_rtl)
        if buf_rtl is None or kind == buf_rtl:
            buf, buf_rtl = buf + ch, kind
        else:
            segments.append((buf, buf_rtl))
            buf, buf_rtl = ch, kind
    if buf:
        segments.append((buf, buf_rtl))
    return segments


def _set_rtl_run(run, font, size, bold, italic, color=None):
    """כותב את כל דגלי ה-rPr הנכונים על run בודד, בסדר הסכמה של OOXML.

    סדר הילדים חייב להישמר: rFonts, b, bCs, i, iCs, color, sz, szCs, rtl.
    """
    rPr = run._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font, qn('w:hAnsi'): font, qn('w:cs'): font}))
    if bold:
        # bold דורש גם w:b (לטינית) וגם w:bCs (כתב מורכב / עברית)
        rPr.append(rPr.makeelement(qn('w:b'), {}))
        rPr.append(rPr.makeelement(qn('w:bCs'), {}))
    if italic:
        rPr.append(rPr.makeelement(qn('w:i'), {}))
        rPr.append(rPr.makeelement(qn('w:iCs'), {}))
    if color:
        rPr.append(rPr.makeelement(qn('w:color'), {qn('w:val'): color}))
    # w:sz/w:szCs בחצאי-נקודה. w:szCs הוא מה שמחיל את הגודל על העברית.
    rPr.append(rPr.makeelement(qn('w:sz'),   {qn('w:val'): str(int(size * 2))}))
    rPr.append(rPr.makeelement(qn('w:szCs'), {qn('w:val'): str(int(size * 2))}))


# סדר הילדים החוקי של CT_PPr לפי סכמת OOXML. חובה לשמור עליו, אחרת Word
# ו-Google Docs מתעלמים מ-w:bidi כשהוא במקום הלא נכון והמסמך נשפך שמאלה.
_PPR_ORDER = [
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
    'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs',
    'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
    'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
    'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
    'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
    'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr',
    'pPrChange',
]


def _insert_ordered(pPr, element):
    """מכניס element ל-pPr במיקום החוקי לפי _PPR_ORDER."""
    tag = element.tag.split('}')[-1]
    idx = _PPR_ORDER.index(tag)
    for child in pPr:
        ctag = child.tag.split('}')[-1]
        if ctag in _PPR_ORDER and _PPR_ORDER.index(ctag) > idx:
            child.addprevious(element)
            return
    pPr.append(element)


def set_para_rtl(paragraph, align='right'):
    """מגדיר בסיס פסקה RTL (w:bidi) ויישור (w:jc) בסדר הסכמה הנכון.

    זה התיקון הקריטי: מסיר bidi/jc קיימים ומכניס אותם מחדש בסדר חוקי, כך
    שגם Word וגם Google Docs מכבדים את ה-RTL. align מקבל right/left/center/
    justify (או ערכי jc גולמיים).
    """
    jc_val = {'right': 'right', 'left': 'left', 'center': 'center',
              'justify': 'both', 'both': 'both'}.get(align, align)
    pPr = paragraph._p.get_or_add_pPr()
    for tag in ('w:bidi', 'w:jc'):
        e = pPr.find(qn(tag))
        if e is not None:
            pPr.remove(e)
    _insert_ordered(pPr, OxmlElement('w:bidi'))
    if jc_val:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), jc_val)
        _insert_ordered(pPr, jc)


def write_rtl_runs(paragraph, text, font, size, bold=False, italic=False,
                   color=None, align='right'):
    """ממלא פסקה קיימת ב-runs מפוצלים-לפי-כתב עם דגלי RTL נכונים.

    מסמן rtl רק על ה-runs העבריים; runs לטיניים נשארים LTR. גם מגדיר את
    בסיס הפסקה RTL והיישור בסדר הסכמה הנכון.
    """
    set_para_rtl(paragraph, align=align)
    for segment, is_rtl in split_by_script(text):
        run = paragraph.add_run(segment)
        _set_rtl_run(run, font, size, bold, italic, color)
        if is_rtl:
            run._r.get_or_add_rPr().append(run._r.rPr.makeelement(qn('w:rtl'), {}))
    return paragraph


def _set_cell_rtl_text(cell, text, font, size, bold=False, color=None,
                       align='right'):
    """מאפס את פסקת ברירת המחדל של תא וממלא אותה בטקסט RTL נכון."""
    cell.text = ''
    p = cell.paragraphs[0]
    write_rtl_runs(p, str(text), font, size, bold=bold, color=color, align=align)


def _shade_cell(cell, fill_hex):
    """צובע רקע של תא טבלה (hex ללא #)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


class HebrewDoc:
    """עטיפה נוחה לבניית מסמך Word עברי מעוצב עם RTL מלא.

    Parameters
    ----------
    font : str   ברירת מחדל לגוף המסמך. בחר גופן שקיים אצל קוראי המסמך:
                 "David" (בטוח ב-Windows+Office), "Assistant"/"Heebo" (מודרני,
                 דורש התקנה), "Frank Ruhl Libre" (סריף למסמכים פורמליים).
    size : int   גודל גוף בנקודות.
    accent : str צבע hex (ללא #) לכותרות ולכותרות טבלה.
    """

    def __init__(self, font="Assistant", size=12, accent="1F4E79",
                 palette=None, tint=None):
        if palette:
            try:
                from palettes import PALETTES
            except ImportError:
                import os as _os
                sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
                from palettes import PALETTES
            p = PALETTES.get(palette, PALETTES["navy"])
            accent = p["accent"].lstrip("#")
            tint = (tint or p["tint"]).lstrip("#")
        self.font = font
        self.size = size
        self.accent = accent.lstrip("#")
        self.zebra = (tint.lstrip("#") if tint else "F2F2F2")
        self.doc = Document()
        # ברירת מחדל לסגנון Normal (משפיע על מה שלא נכתב מפורשות)
        self.doc.styles['Normal'].font.name = font
        self.doc.styles['Normal'].font.size = Pt(size)
        self._make_all_sections_rtl()
        self._set_document_rtl_defaults()

    # ---------- הגדרות עמוד / מקטע ----------
    def _set_document_rtl_defaults(self):
        """מוסיף bidi לברירות המחדל של המסמך ולסגנונות הבסיס.

        חגורה-ושלייקס: כך כל פסקה (גם כזו שלא עברה דרך המתודות) יורשת RTL.
        קריטי במיוחד לייבוא ל-Google Docs, שלפעמים מתעלם מ-bidi פר-פסקה
        ונשען על ברירות המחדל."""
        styles_el = self.doc.styles.element
        # docDefaults > pPrDefault > pPr > bidi
        docDefaults = styles_el.find(qn('w:docDefaults'))
        if docDefaults is not None:
            pPrDefault = docDefaults.find(qn('w:pPrDefault'))
            if pPrDefault is None:
                pPrDefault = OxmlElement('w:pPrDefault')
                docDefaults.insert(0, pPrDefault)
            pPr = pPrDefault.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                pPrDefault.append(pPr)
            if pPr.find(qn('w:bidi')) is None:
                _insert_ordered(pPr, OxmlElement('w:bidi'))
        # סגנונות Normal + כותרות
        for sname in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Title'):
            try:
                st = self.doc.styles[sname]
            except KeyError:
                continue
            pPr = st.element.get_or_add_pPr()
            if pPr.find(qn('w:bidi')) is None:
                _insert_ordered(pPr, OxmlElement('w:bidi'))

    def _make_all_sections_rtl(self):
        for section in self.doc.sections:
            sectPr = section._sectPr
            if sectPr.find(qn('w:bidi')) is None:
                sectPr.append(sectPr.makeelement(qn('w:bidi'), {}))

    def page(self, margin_mm=20, size="A4"):
        """מגדיר שוליים אחידים וגודל עמוד לכל המקטעים."""
        for section in self.doc.sections:
            section.top_margin = Mm(margin_mm)
            section.bottom_margin = Mm(margin_mm)
            section.left_margin = Mm(margin_mm)
            section.right_margin = Mm(margin_mm)
            if size.upper() == "A4":
                section.page_width = Mm(210)
                section.page_height = Mm(297)
        return self

    # ---------- בלוקים ----------
    def heading(self, text, level=1, color=None, size=None, align="right"):
        """כותרת RTL. level=0 לכותרת ראשית (Title)."""
        p = self.doc.add_heading(level=level)
        # מנקה את ה-run האוטומטי שלפעמים נוצר
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        sizes = {0: int(self.size * 2), 1: int(self.size * 1.6),
                 2: int(self.size * 1.35), 3: int(self.size * 1.15)}
        eff_size = size or sizes.get(level, self.size)
        write_rtl_runs(p, text, self.font, eff_size, bold=True,
                       color=color or self.accent, align=align)
        return p

    def paragraph(self, text, bold=False, italic=False, size=None,
                  color=None, align="right", space_after=6):
        p = self.doc.add_paragraph()
        write_rtl_runs(p, text, self.font, size or self.size,
                       bold=bold, italic=italic, color=color, align=align)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullets(self, items, size=None):
        return self._list(items, "List Bullet", size)

    def numbered(self, items, size=None):
        return self._list(items, "List Number", size)

    def _list(self, items, style, size):
        out = []
        for item in items:
            p = self.doc.add_paragraph(style=style)
            write_rtl_runs(p, str(item), self.font, size or self.size,
                           align='right')
            out.append(p)
        return out

    def table(self, headers, rows, col_widths_mm=None, header_fill=None,
              zebra=None, size=None):
        """טבלת RTL: bidiVisual (העמודה הראשונה מימין), כותרת צבועה, פסים."""
        size = size or self.size
        zebra = zebra if zebra is not None else self.zebra
        n = len(headers)
        t = self.doc.add_table(rows=1, cols=n)
        t.style = 'Table Grid'
        t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # bidiVisual: הופך את סדר העמודות כך שהראשונה מימין
        tblPr = t._tbl.tblPr
        tblPr.append(tblPr.makeelement(qn('w:bidiVisual'), {}))

        hdr = t.rows[0].cells
        for i, htext in enumerate(headers):
            _set_cell_rtl_text(hdr[i], htext, self.font, size, bold=True,
                               color="FFFFFF")
            _shade_cell(hdr[i], header_fill or self.accent)

        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            for i in range(n):
                val = row[i] if i < len(row) else ""
                _set_cell_rtl_text(cells[i], val, self.font, size)
                if zebra and ridx % 2 == 1:
                    _shade_cell(cells[i], zebra)

        if col_widths_mm:
            for i, w in enumerate(col_widths_mm):
                for row in t.rows:
                    row.cells[i].width = Mm(w)
        return t

    def spacer(self, pts=10):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pts)
        return p

    def page_break(self):
        from docx.enum.text import WD_BREAK
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- כותרת עליונה/תחתונה עם מספרי עמודים ----------
    def footer_page_numbers(self, prefix="עמוד "):
        for section in self.doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            write_rtl_runs(p, prefix, self.font, self.size - 1, align='center')
            run = p.add_run()
            _set_rtl_run(run, self.font, self.size - 1, False, False)
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), 'PAGE')
            run._r.addnext(fld)
        return self

    def save(self, path):
        self.doc.save(path)
        return path


_ALIGN = {
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ------------------- בדיקה עצמית -------------------
if __name__ == "__main__":
    d = HebrewDoc(font="Assistant", size=12, accent="1F4E79").page(margin_mm=22)
    d.heading("הצעת מחיר", level=0)
    d.paragraph('לכבוד: חברת Acme בע"מ | תאריך: 22/06/2026', color="666666")
    d.heading("פירוט השירותים", level=2)
    d.paragraph(
        'להלן הצעת מחיר לפרויקט הטמעת AI הכולל 3 שלבים: אפיון, פיתוח MVP, '
        'והדרכת צוות. כל השלבים כוללים ליווי שוטף ב-WhatsApp.')
    d.bullets([
        "אפיון מלא של תהליכי העבודה",
        "פיתוח אב-טיפוס (MVP) תוך 30 יום",
        "הדרכת צוות בן 12 עובדים",
    ])
    d.heading("טבלת תמחור", level=2)
    d.table(
        headers=["תיאור", "כמות", 'מחיר ליחידה (ש"ח)', 'סה"כ (ש"ח)'],
        rows=[
            ["אפיון", "1", "5,000.00", "5,000.00"],
            ["פיתוח MVP", "1", "12,000.00", "12,000.00"],
            ["הדרכה", "12", "800.00", "9,600.00"],
        ],
        col_widths_mm=[80, 20, 35, 35],
    )
    d.spacer(8)
    d.paragraph('סה"כ לפני מע"מ: 26,600.00 ש"ח', bold=True, align="right")
    d.paragraph('מע"מ (18%): 4,788.00 ש"ח', align="right")
    d.paragraph('סה"כ לתשלום: 31,388.00 ש"ח', bold=True, size=14,
                color="1F4E79", align="right")
    d.footer_page_numbers()
    out = d.save("/home/claude/test_proposal.docx")
    print("נוצר:", out)
